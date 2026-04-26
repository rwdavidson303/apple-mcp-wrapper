"""Generate a Word doc with Apple Music catalog links for Soul Power tracks
not yet in the 'Soul Power - Suggested Additions' playlist. User clicks each
link to open in Music and adds via the '+' button.
"""
from __future__ import annotations

import asyncio
import subprocess
import urllib.request
from collections import defaultdict

from apple_mcp_wrapper import catalog
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from populate_soul_power import TARGETS

CATEGORY_LABELS = {
    "donny": "Donny Hathaway (foundational male voice #1)",
    "marvin": "Marvin Gaye",
    "stevie": "Stevie Wonder",
    "sam": "Sam Cooke (60s foundation)",
    "otis": "Otis Redding",
    "smokey": "Smokey Robinson & The Miracles",
    "curtis": "Curtis Mayfield / The Impressions",
    "gladys": "Gladys Knight & The Pips (female rebalance)",
    "minnie": "Minnie Riperton",
    "phyllis": "Phyllis Hyman",
    "roberta": "Roberta Flack (solo)",
    "chaka": "Chaka Khan / Rufus",
    "luther": "Luther Vandross (deeper cuts)",
    "stylistics": "The Stylistics (Philly soft-soul)",
    "delfonics": "The Delfonics",
    "bluemagic": "Blue Magic",
    "maining": "The Main Ingredient",
    "intruders": "The Intruders",
    "bill": "Bill Withers",
    "isaac": "Isaac Hayes",
    "spinners": "The Spinners (more cuts)",
    "chilites": "The Chi-Lites (deep cuts)",
    "dells": "The Dells",
    "modern": "Modern torch-bearers",
}


def url_ok(url: str) -> bool:
    req = urllib.request.Request(url, method="HEAD")
    try:
        resp = urllib.request.urlopen(req, timeout=8)
        return 200 <= resp.status < 400
    except Exception:
        return False


def playlist_track_keys() -> set[tuple[str, str]]:
    """Return a set of (lower_title, lower_artist) pairs already in the playlist."""
    script = '''
    set outStr to ""
    tell application "Music"
        try
            set pl to user playlist "Soul Power - Suggested Additions"
            repeat with t in (every track of pl)
                set outStr to outStr & (name of t) & "||" & (artist of t) & "\n"
            end repeat
        end try
    end tell
    return outStr
    '''
    result = subprocess.run(
        ["osascript", "-e", script],
        capture_output=True, text=True, timeout=15,
    )
    keys = set()
    for line in result.stdout.splitlines():
        if "||" in line:
            name, artist = line.split("||", 1)
            keys.add((name.strip().lower(), artist.strip().lower()))
    return keys


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


async def main() -> None:
    have = playlist_track_keys()
    print(f"Already in playlist: {len(have)} tracks")

    doc = Document()
    doc.add_heading("Soul Power - Suggested Additions (remaining)", 0)
    intro = doc.add_paragraph()
    intro.add_run(
        "Click each link to open the track in Apple Music. Tap '+' to add to "
        "your library. When you are done, tell Claude and the tracks will be "
        "pulled into the 'Soul Power - Suggested Additions' playlist."
    ).italic = True

    remaining_by_category: dict[str, list[tuple[str, str, dict]]] = defaultdict(list)
    skip_count = 0
    stale_count = 0
    no_match_count = 0

    for (category, artist, title) in TARGETS:
        for attempt in range(3):
            try:
                r = await catalog.find_best_match(artist, title)
                break
            except Exception as e:
                if "429" in str(e) and attempt < 2:
                    await asyncio.sleep(10)
                    continue
                print(f"search-err: {artist} - {title}: {e}")
                r = None
                break
        await asyncio.sleep(0.5)
        if not r:
            no_match_count += 1
            print(f"no-match: {artist} - {title}")
            continue

        resolved_title = r.get("trackName", "")
        resolved_artist = r.get("artistName", "")
        key = (resolved_title.lower(), resolved_artist.lower())
        if key in have:
            skip_count += 1
            continue

        url = catalog.canonical_url(r)
        if not url or not url_ok(url):
            stale_count += 1
            print(f"stale: {artist} - {title}")
            continue

        remaining_by_category[category].append((artist, title, r))

    total_remaining = sum(len(v) for v in remaining_by_category.values())
    print(f"remaining to add: {total_remaining}, already-in-playlist: {skip_count}, "
          f"stale-url: {stale_count}, no-match: {no_match_count}")

    for category_key in CATEGORY_LABELS:
        items = remaining_by_category.get(category_key, [])
        if not items:
            continue
        doc.add_heading(CATEGORY_LABELS[category_key], level=1)
        for (artist, title, r) in items:
            url = catalog.canonical_url(r)
            label = f"{r['trackName']} - {r['artistName']} ({r.get('collectionName','')})"
            p = doc.add_paragraph(style="List Bullet")
            add_hyperlink(p, url, label)

    out = "/Users/richarddavidson/Desktop/Desktop - Mac/Claude/Soul_Power_Additions_Links.docx"
    doc.save(out)
    print(f"\nSaved: {out}")
    subprocess.run(["open", "-a", "Microsoft Word", out])


if __name__ == "__main__":
    asyncio.run(main())
