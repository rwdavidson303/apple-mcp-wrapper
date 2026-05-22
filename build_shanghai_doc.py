"""Resolve a curated list of additions to Nyah's "shanghai" Spotify playlist
against the Apple Music catalog (sanity check), then build a Word doc with
one-click Spotify search URLs per track for manual hand-off.
"""
from __future__ import annotations

import asyncio
import subprocess
import urllib.parse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from apple_mcp_wrapper import catalog


SECTIONS: list[tuple[str, list[tuple[str, str]]]] = [
    ("SAULT / Cleo Sol extensions", [
        ("Cleo Sol", "Why Don't You"),
        ("SAULT", "Wildfires"),
        ("SAULT", "Free"),
    ]),
    ("Charlotte Day Wilson lane", [
        ("Charlotte Day Wilson", "Mountains"),
        ("Snoh Aalegra", "I Want You Around"),
        ("Yazmin Lacey", "Bad Company"),
    ]),
    ("Daniel Caesar / Toronto-soul", [
        ("Daniel Caesar", "Best Part (feat. H.E.R.)"),
        ("Mustafa", "Stay Alive"),
        ("Sampha", "(No One Knows Me) Like the Piano"),
    ]),
    ("UK contemporary soul", [
        ("Jorja Smith", "Be Honest"),
        ("Joy Crookes", "Feet Don't Fail Me Now"),
        ("Greentea Peng", "Hu Man"),
        ("Lianne La Havas", "Bittersweet"),
    ]),
    ("Frank Ocean / Kali Uchis lane", [
        ("Frank Ocean", "White Ferrari"),
        ("Frank Ocean", "Self Control"),
        ("Kali Uchis", "After the Storm"),
        ("Solange", "Cranes in the Sky"),
    ]),
    ("Sade canon", [
        ("Sade", "By Your Side"),
        ("Sade", "Cherish the Day"),
    ]),
    ("Indie folk to pair with Big Thief", [
        ("Phoebe Bridgers", "Motion Sickness"),
        ("Sufjan Stevens", "Mystery of Love"),
        ("Adrianne Lenker", "anything"),
    ]),
    ("Electronic downtempo", [
        ("James Blake", "Retrograde"),
        ("The xx", "Angels"),
        ("Jamie xx", "Loud Places (feat. Romy)"),
    ]),
    ("Vintage soul anchors", [
        ("Bill Withers", "Ain't No Sunshine"),
        ("Nina Simone", "Feeling Good"),
        ("Roberta Flack", "The First Time Ever I Saw Your Face"),
    ]),
    ("Maribou State / Bonobo lane", [
        ("Maribou State", "Slow Heat"),
        ("Bonobo", "Kerala"),
    ]),
]


def spotify_search_url(artist: str, title: str) -> str:
    """Build a Spotify search URL that opens to a pre-filled search."""
    query = f"{artist} {title}"
    return f"https://open.spotify.com/search/{urllib.parse.quote(query)}"


def add_hyperlink(paragraph, url: str, text: str):
    """Add a clickable hyperlink to a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


async def resolve_all() -> list[dict]:
    """Resolve every track against Apple Music to verify it exists.
    Returns a list of dicts with section, artist, title, resolved, url."""
    results: list[dict] = []
    for section, tracks in SECTIONS:
        for artist, title in tracks:
            try:
                m = await catalog.find_best_match(artist, title)
            except Exception as e:
                m = None
                print(f"  exception on {artist} - {title}: {e}")
            if m and m.get("trackId"):
                resolved = f'{m.get("artistName", "")} - {m.get("trackName", "")}'
                ok = True
            else:
                # Fallback to plain iTunes search
                hits = await catalog.search(f"{artist} {title}", limit=5)
                if hits:
                    h = hits[0]
                    resolved = f'{h.get("artistName", "")} - {h.get("trackName", "")}'
                    ok = True
                else:
                    resolved = "NOT FOUND on Apple Music catalog"
                    ok = False
            results.append({
                "section": section,
                "artist": artist,
                "title": title,
                "resolved": resolved,
                "ok": ok,
                "url": spotify_search_url(artist, title),
            })
            print(f"[{'OK ' if ok else 'MISS'}] {artist} - {title}  ->  {resolved}")
            await asyncio.sleep(0.2)
    return results


def build_docx(results: list[dict], out_path: Path) -> None:
    doc = Document()

    h = doc.add_heading("shanghai — playlist enhancements", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)

    intro = doc.add_paragraph()
    intro.add_run(
        "Thirty tracks that match the vibe of your shanghai playlist — "
        "modern neo-soul, downtempo, intimate, in the SAULT / Cleo Sol / "
        "Charlotte Day Wilson lane with a couple of vintage anchors. "
        "Tap any track name and Spotify will open to a pre-filled search; "
        "tap the result, then tap the three-dot menu and "
        "“Add to playlist.”"
    )

    last_section = None
    track_num = 0
    for r in results:
        if r["section"] != last_section:
            last_section = r["section"]
            doc.add_paragraph()
            sh = doc.add_paragraph()
            sh_run = sh.add_run(r["section"])
            sh_run.bold = True
            sh_run.font.size = Pt(12)
        track_num += 1
        p = doc.add_paragraph(style="List Number")
        link_text = f'{r["artist"]} — {r["title"]}'
        add_hyperlink(p, r["url"], link_text)
        if not r["ok"]:
            p.add_run("  (verify manually)").italic = True

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot_run = foot.add_run(
        "Built with love by Richard. xo"
    )
    foot_run.italic = True
    foot_run.font.size = Pt(10)

    doc.save(out_path)


def main() -> None:
    out_dir = Path.home() / "Desktop" / "Desktop - Mac" / "Claude" / "Shanghai-Playlist"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "shanghai-playlist-enhancements.docx"

    print("Resolving 30 tracks against Apple Music catalog...\n")
    results = asyncio.run(resolve_all())

    n_ok = sum(1 for r in results if r["ok"])
    n_miss = len(results) - n_ok
    print(f"\nResolved: {n_ok}/{len(results)}  (misses: {n_miss})")
    if n_miss:
        print("Misses:")
        for r in results:
            if not r["ok"]:
                print(f"  - {r['artist']} - {r['title']}")

    print(f"\nWriting Word doc to {out_path}")
    build_docx(results, out_path)

    print("Opening in Microsoft Word...")
    subprocess.run(["open", "-a", "Microsoft Word", str(out_path)], check=False)
    print("Done.")


if __name__ == "__main__":
    main()
